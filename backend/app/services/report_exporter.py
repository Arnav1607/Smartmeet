import io
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def generate_excel_report(meeting, report, participants, tasks) -> io.BytesIO:
    """Generate a highly styled, professional Excel workbook containing all meeting insights."""
    wb = Workbook()
    
    # Setup Sheet 1: Overview & Summary
    ws1 = wb.active
    ws1.title = "Executive Summary"
    ws1.views.sheetView[0].showGridLines = True
    
    # Theme colors
    HEADER_FILL = PatternFill(start_color="1A1A2E", end_color="1A1A2E", fill_type="solid")
    ACCENT_FILL = PatternFill(start_color="7C6FFF", end_color="7C6FFF", fill_type="solid")
    ZEBRA_FILL = PatternFill(start_color="F5F5F8", end_color="F5F5F8", fill_type="solid")
    WHITE_FILL = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
    
    FONT_TITLE = Font(name="Segoe UI", size=16, bold=True, color="FFFFFF")
    FONT_SECTION = Font(name="Segoe UI", size=12, bold=True, color="1A1A2E")
    FONT_HEADER = Font(name="Segoe UI", size=10, bold=True, color="FFFFFF")
    FONT_BODY_BOLD = Font(name="Segoe UI", size=10, bold=True, color="000000")
    FONT_BODY = Font(name="Segoe UI", size=10, color="333333")
    
    # Borders
    thin = Side(border_style="thin", color="D3D3D3")
    border_all = Border(left=thin, right=thin, top=thin, bottom=thin)
    
    # Title Block
    ws1.merge_cells("A1:D2")
    title_cell = ws1["A1"]
    title_cell.value = "SmartMeet AI — Meeting Intelligence Brief"
    title_cell.font = FONT_TITLE
    title_cell.fill = HEADER_FILL
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    
    # Set rows height
    ws1.row_dimensions[1].height = 25
    ws1.row_dimensions[2].height = 25
    
    # Metadata block
    ws1["A4"] = "Meeting Details"
    ws1["A4"].font = FONT_SECTION
    
    details = [
        ("Meeting Title", meeting.title),
        ("Platform", meeting.platform.upper()),
        ("Date", meeting.started_at.strftime('%Y-%m-%d %H:%M UTC') if meeting.started_at else "—"),
        ("Duration", f"{meeting.duration_mins} minutes"),
        ("Total Participants", len(participants)),
        ("Productivity Score", f"{report.productivity_score}/100" if report else "—"),
        ("Average Sentiment", f"{report.sentiment_score} ({'Positive' if (report and report.sentiment_score > 0.1) else 'Negative' if (report and report.sentiment_score < -0.1) else 'Neutral'})" if report else "—")
    ]
    
    curr_row = 5
    for label, val in details:
        ws1.cell(row=curr_row, column=1, value=label).font = FONT_BODY_BOLD
        ws1.cell(row=curr_row, column=1).fill = ZEBRA_FILL
        ws1.cell(row=curr_row, column=1).border = border_all
        
        ws1.cell(row=curr_row, column=2, value=val).font = FONT_BODY
        ws1.cell(row=curr_row, column=2).border = border_all
        ws1.merge_cells(start_row=curr_row, start_column=2, end_row=curr_row, end_column=4)
        curr_row += 1
        
    # Executive Summary Text
    curr_row += 1
    ws1.cell(row=curr_row, column=1, value="Executive Summary").font = FONT_SECTION
    curr_row += 1
    ws1.merge_cells(start_row=curr_row, start_column=1, end_row=curr_row+4, end_column=4)
    summary_cell = ws1.cell(row=curr_row, column=1, value=report.executive_summary if report else "Processing summary...")
    summary_cell.font = FONT_BODY
    summary_cell.alignment = Alignment(wrap_text=True, vertical="top")
    for r in range(curr_row, curr_row+5):
        for c in range(1, 5):
            ws1.cell(row=r, column=c).border = border_all
            
    # Sheet 2: Action Items & Tasks
    ws2 = wb.create_sheet(title="Action Items")
    ws2.views.sheetView[0].showGridLines = True
    ws2.row_dimensions[1].height = 25
    
    headers = ["Task ID", "Task Description", "Assigned Owner", "Deadline", "Priority", "Status"]
    for col_idx, text in enumerate(headers, 1):
        cell = ws2.cell(row=1, column=col_idx, value=text)
        cell.font = FONT_HEADER
        cell.fill = ACCENT_FILL
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border_all
        
    for r_idx, task in enumerate(tasks, 2):
        ws2.row_dimensions[r_idx].height = 20
        row_data = [
            task.task_id[:8],
            task.description,
            task.owner_name or "Unassigned",
            task.deadline.strftime('%Y-%m-%d') if task.deadline else "—",
            task.priority.upper(),
            task.status.upper()
        ]
        for c_idx, val in enumerate(row_data, 1):
            cell = ws2.cell(row=r_idx, column=c_idx, value=val)
            cell.font = FONT_BODY
            cell.border = border_all
            cell.fill = ZEBRA_FILL if r_idx % 2 == 0 else WHITE_FILL
            
            # Align center for fields that look better
            if c_idx in [1, 4, 5, 6]:
                cell.alignment = Alignment(horizontal="center")
                
    # Sheet 3: Attendance & Speakers
    ws3 = wb.create_sheet(title="Attendance & Speaker Stats")
    ws3.views.sheetView[0].showGridLines = True
    
    # Participants Title
    ws3["A1"] = "Participant Attendance Logs"
    ws3["A1"].font = FONT_SECTION
    ws3.row_dimensions[2].height = 25
    
    att_headers = ["Name", "Email Address", "Joined At", "Left At", "Speaking Minutes"]
    for col_idx, text in enumerate(att_headers, 1):
        cell = ws3.cell(row=2, column=col_idx, value=text)
        cell.font = FONT_HEADER
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border_all
        
    row_cursor = 3
    for p in participants:
        ws3.row_dimensions[row_cursor].height = 20
        joined_str = p.joined_at.strftime('%H:%M:%S') if p.joined_at else "—"
        left_str = p.left_at.strftime('%H:%M:%S') if p.left_at else "—"
        
        row_data = [p.name, p.email, joined_str, left_str, f"{p.duration_mins} mins"]
        for col_idx, val in enumerate(row_data, 1):
            cell = ws3.cell(row=row_cursor, column=col_idx, value=val)
            cell.font = FONT_BODY
            cell.border = border_all
            cell.fill = ZEBRA_FILL if row_cursor % 2 == 0 else WHITE_FILL
            if col_idx in [3, 4, 5]:
                cell.alignment = Alignment(horizontal="center")
        row_cursor += 1
        
    # Speaker stats subtable
    row_cursor += 2
    ws3.cell(row=row_cursor, column=1, value="Speaker Metrics").font = FONT_SECTION
    row_cursor += 1
    
    spk_headers = ["Speaker Name", "Total Messages", "Word Count", "Participation %"]
    for col_idx, text in enumerate(spk_headers, 1):
        cell = ws3.cell(row=row_cursor, column=col_idx, value=text)
        cell.font = FONT_HEADER
        cell.fill = ACCENT_FILL
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border_all
    
    row_cursor += 1
    if report and report.speaker_stats:
        for name, stats in report.speaker_stats.items():
            row_data = [
                name, 
                stats.get('messages', 0), 
                stats.get('words', 0), 
                f"{stats.get('participation_pct', 0)}%"
            ]
            for col_idx, val in enumerate(row_data, 1):
                cell = ws3.cell(row=row_cursor, column=col_idx, value=val)
                cell.font = FONT_BODY
                cell.border = border_all
                if col_idx in [2, 3, 4]:
                    cell.alignment = Alignment(horizontal="center")
            row_cursor += 1
    else:
        ws3.cell(row=row_cursor, column=1, value="Speaker statistics not available.").font = FONT_BODY
        ws3.merge_cells(start_row=row_cursor, start_column=1, end_row=row_cursor, end_column=4)

    # Auto-adjust column widths for all sheets
    for sheet in wb.worksheets:
        for col in sheet.columns:
            max_len = 0
            col_letter = get_column_letter(col[0].column)
            for cell in col:
                val_str = str(cell.value or '')
                # Skip merged cells or long summary briefs when checking size
                if len(val_str) > 50:
                    continue
                if len(val_str) > max_len:
                    max_len = len(val_str)
            sheet.column_dimensions[col_letter].width = max(max_len + 4, 12)
            
    # Save to buffer
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer

def generate_docx_report(meeting, report, participants, tasks) -> io.BytesIO:
    """Generate a high-quality Microsoft Word (.docx) document as an Executive Meeting Brief."""
    doc = Document()
    
    # Document Styling Page Width Setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set default fonts
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Helper function to color cells
    def set_cell_background(cell, color_hex):
        shading_xml = f'<w:shd {qn("w:fill")}="{color_hex}"/>'
        cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd'))
        cell._tc.get_or_add_tcPr().last_child.set(qn('w:fill'), color_hex)

    # Title Banner
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = p_title.add_run("SMARTMEET AI")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x7C, 0x6F, 0xFF) # Brand Purple

    p_subtitle = doc.add_paragraph()
    run_sub = p_subtitle.add_run("Intelligent Meeting Assistant & Summary Report")
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Separator Line
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_after = Pt(20)
    p_sep_run = p_sep.add_run("―" * 50)
    p_sep_run.font.color.rgb = RGBColor(0xD3, 0xD3, 0xD3)

    # Metadata Grid
    doc.add_heading("Meeting Briefing Information", level=2)
    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.style = 'Table Grid'
    
    meta_fields = [
        ("Meeting Title", meeting.title),
        ("Platform Source", meeting.platform.upper()),
        ("Start Date & Time", meeting.started_at.strftime('%B %d, %Y - %H:%M UTC') if meeting.started_at else "—"),
        ("Meeting Length", f"{meeting.duration_mins} minutes"),
        ("Overall Team Productivity", f"{report.productivity_score}/100" if report else "—"),
        ("Meeting Type Tag", report.meeting_type.upper() if report else "STANDUP")
    ]
    
    for idx, (label, val) in enumerate(meta_fields):
        row = meta_table.rows[idx]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        set_cell_background(row.cells[0], "F5F5F8")
        row.cells[1].text = str(val)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Executive Summary Section
    doc.add_heading("1. Executive Summary", level=2)
    p_exec = doc.add_paragraph()
    p_exec.paragraph_format.line_spacing = 1.25
    p_exec.add_run(report.executive_summary if report else "No summary available.")
    
    # Detailed Summary Section
    if report and report.detailed_summary:
        doc.add_heading("2. Detailed Discussions", level=2)
        p_detail = doc.add_paragraph()
        p_detail.paragraph_format.line_spacing = 1.25
        p_detail.add_run(report.detailed_summary)

    # Key Decisions
    doc.add_heading("3. Core Decisions Logged", level=2)
    decisions = report.key_decisions if report else []
    if decisions:
        for d in decisions:
            doc.add_paragraph(d, style='List Bullet')
    else:
        doc.add_paragraph("No specific key decisions were logged.", style='Normal')

    # Action Items Table
    doc.add_heading("4. Extracted Action Items", level=2)
    if tasks:
        task_table = doc.add_table(rows=1, cols=4)
        task_table.style = 'Table Grid'
        
        # Headers
        headers = ["Task Description", "Assigned Owner", "Deadline", "Priority"]
        for idx, text in enumerate(headers):
            cell = task_table.rows[0].cells[idx]
            cell.text = text
            cell.paragraphs[0].runs[0].font.bold = True
            set_cell_background(cell, "7C6FFF")
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
        for t in tasks:
            row = task_table.add_row()
            row.cells[0].text = t.description
            row.cells[1].text = t.owner_name or "Unassigned"
            row.cells[2].text = t.deadline.strftime('%Y-%m-%d') if t.deadline else "—"
            row.cells[3].text = t.priority.upper()
    else:
        doc.add_paragraph("No action items were assigned in this session.", style='Normal')

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Attendance logs
    doc.add_heading("5. Participant List & Speaking Attendance", level=2)
    if participants:
        part_table = doc.add_table(rows=1, cols=3)
        part_table.style = 'Table Grid'
        
        headers = ["Participant Name", "Email Address", "Duration Active"]
        for idx, text in enumerate(headers):
            cell = part_table.rows[0].cells[idx]
            cell.text = text
            cell.paragraphs[0].runs[0].font.bold = True
            set_cell_background(cell, "1A1A2E")
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
        for p in participants:
            row = part_table.add_row()
            row.cells[0].text = p.name
            row.cells[1].text = p.email or "—"
            row.cells[2].text = f"{p.duration_mins} minutes"
    else:
        doc.add_paragraph("No attendee information logged.", style='Normal')

    # Footer signature
    doc.add_paragraph().paragraph_format.space_after = Pt(30)
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = p_footer.add_run(f"Briefing automatically compiled by SmartMeet AI\nDate generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}")
    footer_run.font.size = Pt(8.5)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save to stream
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
